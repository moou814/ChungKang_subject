from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUT = ROOT / "outputs" / "ChungKang_subject_프로그램_기술명세서.docx"

FONT = "Malgun Gothic"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 44)
MUTED = RGBColor(80, 80, 80)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=None, name=FONT, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(p, before=0, after=6, line=1.10):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="B8BCC4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_GRAY, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_text(hdr[i], text, bold=True, fill=header_fill, size=font_size)
    for row_data in rows:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            set_cell_text(row[i], text, size=font_size)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=4)
    return table


def set_cell_text(cell, text, bold=False, fill=None, size=9.5, color=INK, align=None):
    if fill:
        set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=0, line=1.10)
    if p.runs:
        p.runs[0].text = ""
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def para(doc, text="", size=11, bold=False, color=INK, style=None, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, before=before, after=after)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, bold=True, color=BLUE)
        set_paragraph_spacing(p, before=16, after=8, line=1.10)
    elif level == 2:
        set_run_font(r, size=13, bold=True, color=BLUE)
        set_paragraph_spacing(p, before=12, after=6, line=1.10)
    else:
        set_run_font(r, size=12, bold=True, color=DARK_BLUE)
        set_paragraph_spacing(p, before=8, after=4, line=1.10)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph_spacing(p, after=4, line=1.15)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def code_block(doc, lines):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=8, line=1.05)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F7F9FB")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), "D7DBE2")
        borders.append(node)
    p_pr.append(borders)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=8.7, color=RGBColor(35, 45, 60))


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name in ("Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    hdr = section.header.paragraphs[0]
    hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hdr.add_run("2026 청강게임대전 게임프로그래밍 | 프로그램 기술 명세서")
    set_run_font(r, size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_cover(doc):
    para(doc, "2026 청강게임대전 게임프로그래밍", size=11, bold=True, color=MUTED, after=2)
    p = para(doc, "프로그램 기술 명세서", size=26, bold=True, color=RGBColor(0, 0, 0), after=3)
    p.paragraph_format.space_before = Pt(0)
    para(doc, "Unity 기반 수학/자료구조 퍼즐 게임 ChungKang_subject", size=14, color=RGBColor(55, 55, 55), after=14)

    meta_rows = [
        ("작품/프로젝트명", "CK_project / ChungKang_subject"),
        ("분야", "게임프로그래밍"),
        ("작성일", "2026년 7월 9일"),
        ("작성 기준", "작업 폴더 내 Unity 프로젝트 코드, ScriptableObject 스테이지 데이터, 대회 과제 안내서"),
        ("제출자 정보", "확인 필요: 성명, 생년월일, 팀명 또는 개인 참가 여부"),
    ]
    add_table(doc, ["항목", "내용"], meta_rows, [2100, 7260], header_fill=LIGHT_BLUE, font_size=10)

    para(
        doc,
        "본 문서는 제출 안내서의 프로그램 기술 명세서 요구사항에 맞춰 핵심 알고리즘, 클래스 구조, "
        "개발 환경, 외부 리소스 및 AI 활용 내역을 정리한 제출용 기술 문서입니다.",
        size=10.5,
        color=INK,
        after=10,
    )


def add_requirement_mapping(doc):
    heading(doc, "1. 대회 요구사항 대응 요약", 1)
    rows = [
        ("수학/자료구조 기반 퍼즐", "비트 연산, 벡터 반사, 2차원 그리드와 DFS 탐색을 핵심 퍼즐 판정에 사용", "충족"),
        ("3개 이상 레벨", "파이프/비트마스크/거울 퍼즐 각각 ScriptableObject 기반 3개 스테이지 보유", "충족"),
        ("입력-연산-결과 검증", "스위치 클릭, 블록 회전, 거울 회전 입력 후 즉시 상태 재계산 및 Clear 판정", "충족"),
        ("게임 상태 전환", "MainScene에서 퍼즐 선택 후 PipeScene, BitMaskScene, MirrorScene으로 전환", "충족"),
        ("디버그 콘솔/치트", "FlowManager.WriteLog로 내부 연산 로그 출력. 레벨 강제 스킵 UI 연결은 제출 전 확인 필요", "부분 충족"),
        ("외부 리소스 명시", "Unity 패키지, TextMesh Pro 리소스, 자체 파이프 이미지 리소스 사용 내역 표기", "문서 반영"),
        ("AI 활용 명시", "본 명세서 작성 및 코드 분석 보조 내역을 AI 활용 명세표에 표기", "문서 반영"),
    ]
    add_table(doc, ["요구 항목", "프로젝트 대응 내용", "판정"], rows, [1900, 5960, 1500], header_fill=LIGHT_BLUE, font_size=9.2)


def add_overview(doc):
    heading(doc, "2. 게임 개요", 1)
    para(
        doc,
        "ChungKang_subject는 Unity 2D 환경에서 여러 논리 퍼즐을 해결하는 게임입니다. "
        "플레이어는 스테이지를 선택한 뒤 세 종류의 퍼즐을 순서와 관계없이 해결하며, 각 퍼즐은 "
        "입력 직후 내부 상태를 계산하고 클리어 조건을 검증합니다.",
    )
    add_table(
        doc,
        ["구분", "내용"],
        [
            ("장르", "수학적 사고와 자료구조 기반의 2D 논리 퍼즐"),
            ("플레이 목표", "선택한 스테이지 안에서 파이프, 비트마스크, 거울 퍼즐을 모두 클리어"),
            ("핵심 재미", "플레이어의 입력이 내부 상태 배열, 비트 상태, 광선 경로로 바뀌고 결과가 즉시 피드백되는 구조"),
            ("주요 조작", "파이프 블록 클릭 회전, 스위치 클릭, 마우스 좌/우 버튼을 이용한 거울 회전"),
            ("권장 플레이 분량", "1회 플레이 5~10분 내외를 목표로 설계 가능한 분량"),
        ],
        [2100, 7260],
        header_fill=LIGHT_GRAY,
        font_size=9.5,
    )


def add_environment(doc):
    heading(doc, "3. 개발 환경 및 프로젝트 구성", 1)
    add_table(
        doc,
        ["항목", "확인된 값"],
        [
            ("Unity Editor", "6000.2.14f1"),
            ("프로젝트명", "CK_project"),
            ("렌더링", "Universal Render Pipeline 17.2.0, 2D Renderer 설정"),
            ("UI", "Unity UGUI 2.0.0, TextMesh Pro 리소스"),
            ("입력", "Unity Input System 1.16.0, Mouse.current 기반 입력 처리"),
            ("주요 씬", "MainScene, PipeScene, BitMaskScene, MirrorScene"),
            ("타겟 플랫폼", "Standalone 설정 확인. 최종 제출 빌드 타겟은 Unity Build Settings에서 확인 필요"),
        ],
        [2300, 7060],
        header_fill=LIGHT_BLUE,
        font_size=9.5,
    )
    add_table(
        doc,
        ["폴더/파일", "역할"],
        [
            ("Assets/Script/FlowManager.cs", "씬 전환, 스테이지 선택, 퍼즐 클리어 상태, 디버그 로그 관리"),
            ("Assets/Script/pipePuzzle", "파이프 블록 회전, 연결 가능 방향 계산, DFS 기반 도달성 검증"),
            ("Assets/Script/BitMask", "스위치별 비트 패턴과 XOR/OR/AND 연산을 통한 램프 상태 검증"),
            ("Assets/Script/mirrorPuzzle", "거울 회전, Raycast2D 광선 추적, 반사/프리즘/타깃 판정"),
            ("Assets/Resources/*/stage", "ScriptableObject 형태의 스테이지 데이터"),
            ("Assets/Resources/pipePuzzle/image", "파이프 블록 표시용 PNG 이미지"),
        ],
        [3500, 5860],
        header_fill=LIGHT_GRAY,
        font_size=9.0,
    )


def add_scene_flow(doc):
    heading(doc, "4. 게임 루프 및 상태 관리", 1)
    para(
        doc,
        "전체 흐름은 FlowManager가 담당합니다. FlowManager는 DontDestroyOnLoad로 유지되며, "
        "스테이지 선택값과 각 퍼즐 클리어 여부를 보관하고 씬 전환 요청을 수행합니다.",
    )
    code_block(
        doc,
        [
            "MainScene",
            "  -> selectStage(stage)",
            "  -> GoPuzzle(puzzleIndex)",
            "  -> RequestScene(PipeScene | BitMaskScene | MirrorScene)",
            "  -> 각 퍼즐 매니저가 입력마다 Clear 조건 검사",
            "  -> FlowManager.Clear()",
            "  -> 세 퍼즐이 모두 clear이면 스테이지 clear 처리",
        ],
    )
    add_table(
        doc,
        ["상태/함수", "설명", "검토 의견"],
        [
            ("stage", "선택한 스테이지 번호를 저장하고 각 퍼즐 매니저가 이를 참조", "ScriptableObject 배열 인덱스와 일치 여부 확인 필요"),
            ("IsPuzzleClear[3]", "파이프, 비트마스크, 거울 퍼즐 클리어 여부 저장", "퍼즐 재진입 방지에 사용"),
            ("RequestScene", "정수 인덱스를 씬 이름 배열에 매핑하여 LoadScene 호출", "씬 이름과 Build Settings가 일치해야 함"),
            ("WriteLog", "TextMesh Pro UI에 내부 연산 로그를 12줄 단위로 출력", "심사용 디버그 콘솔 요구사항에 대응"),
            ("Clear", "현재 퍼즐 클리어 처리 및 전체 스테이지 클리어 검사", "stage 값이 1~3일 때 IsStageClear 인덱스 보정 필요"),
        ],
        [1900, 4460, 3000],
        header_fill=LIGHT_BLUE,
        font_size=8.8,
    )


def add_bitmask(doc):
    heading(doc, "5. 핵심 알고리즘 1 - 비트마스크 퍼즐", 1)
    para(
        doc,
        "비트마스크 퍼즐은 5개의 램프 상태를 하나의 uint 값으로 압축하여 관리합니다. "
        "각 스위치는 5비트 패턴과 연산 타입(XOR, OR, AND)을 가지고 있으며, 플레이어가 스위치를 누르면 "
        "현재 상태 curState에 비트 연산을 적용합니다. 모든 램프 비트가 1이면 클리어입니다.",
    )
    code_block(
        doc,
        [
            "switch 입력:",
            "  XOR -> curState = curState ^ switchMask",
            "  OR  -> curState = curState | switchMask",
            "  AND -> curState = curState & switchMask",
            "램프 검사:",
            "  for i in 0..4:",
            "    lamp[i] = (curState & (1 << i)) != 0",
            "  모든 lamp[i]가 true이면 Clear",
        ],
    )
    add_table(
        doc,
        ["스테이지", "스위치 타입", "스위치 비트 패턴", "난이도 근거"],
        [
            ("1", "XOR, XOR, XOR, XOR", "10101 / 11000 / 00011 / 01110", "기본 XOR 토글 규칙 중심"),
            ("2", "XOR, OR, XOR", "11001 / 01010 / 01110", "OR 연산 추가로 되돌릴 수 없는 켜짐 상태 고려"),
            ("3", "AND, XOR, XOR", "01001 / 11010 / 01100", "AND 연산으로 기존 상태가 제거될 수 있어 순서 판단 필요"),
        ],
        [1100, 2400, 3100, 2760],
        header_fill=LIGHT_GRAY,
        font_size=8.8,
    )
    add_table(
        doc,
        ["평가 관점", "설명"],
        [
            ("수학/논리 개념", "이진수 비트, 논리 연산자 XOR/OR/AND, 마스크 기반 상태 압축"),
            ("자료구조", "uint 상태값, switchs 배열, Image 배열"),
            ("시간 복잡도", "스위치 입력 1회당 O(5). 램프 수가 고정되어 사실상 상수 시간"),
            ("장점", "램프 상태를 bool 배열 전체가 아닌 하나의 정수로 표현하여 연산이 단순하고 빠름"),
        ],
        [1900, 7460],
        header_fill=LIGHT_BLUE,
        font_size=9.2,
    )


def add_pipe(doc):
    heading(doc, "6. 핵심 알고리즘 2 - 파이프 연결 퍼즐", 1)
    para(
        doc,
        "파이프 퍼즐은 2차원 그리드 위의 블록을 회전시켜 시작점에서 목적지까지 연결되는지를 검사합니다. "
        "각 블록은 모양(L, I, plus, T)과 회전 각도에 따라 이동 가능한 방향을 반환하며, 매 입력 후 DFS 방식으로 "
        "도달 가능한 모든 칸을 탐색합니다.",
    )
    code_block(
        doc,
        [
            "clearCheck():",
            "  visited = false로 초기화",
            "  stack에 start 좌표 push",
            "  while stack이 비어 있지 않음:",
            "    cur = stack.pop()",
            "    cur 블록의 연결 방향 목록 확인",
            "    next가 범위 안이고 아직 방문 전이며",
            "    next 블록이 반대 방향 연결을 가지고 있으면 방문 처리",
            "  모든 목적지 end가 visited이면 Clear",
        ],
    )
    add_table(
        doc,
        ["스테이지", "그리드", "목적지 수", "고정 블록 수", "난이도 근거"],
        [
            ("1", "4 x 6", "1", "0", "기본 연결 규칙과 단일 목적지"),
            ("2", "5 x 6", "1", "5", "고정 블록으로 회전 가능한 선택지 제한"),
            ("3", "5 x 8", "5", "5", "다중 목적지와 더 넓은 연결 경로 검증"),
        ],
        [1000, 1350, 1350, 1450, 4210],
        header_fill=LIGHT_GRAY,
        font_size=8.8,
    )
    add_table(
        doc,
        ["평가 관점", "설명"],
        [
            ("자료구조", "blockKind[,] map, pipePuzzle_Block[,] blocks, bool[,] visited, List<int[]> stack"),
            ("알고리즘", "명시적 스택을 사용하는 DFS 기반 도달성 탐색"),
            ("시간 복잡도", "V=행x열, E는 각 칸 최대 4방향일 때 O(V+E), 즉 O(V)에 가까움"),
            ("실시간 검증", "블록 회전 직후 clearCheck를 호출하여 현재 연결 상태를 즉시 시각화"),
        ],
        [1900, 7460],
        header_fill=LIGHT_BLUE,
        font_size=9.2,
    )


def add_mirror(doc):
    heading(doc, "7. 핵심 알고리즘 3 - 거울/광선 퍼즐", 1)
    para(
        doc,
        "거울 퍼즐은 벡터와 기하학을 활용합니다. Raycast2D로 광선이 충돌하는 오브젝트를 찾고, "
        "거울에 닿으면 입사 방향과 거울의 법선 벡터를 이용해 반사 방향을 계산합니다. 프리즘에 닿으면 "
        "광선을 세 방향으로 분기하도록 설계되어 있습니다.",
    )
    code_block(
        doc,
        [
            "반사 방향 계산:",
            "  dir = 입사 방향 정규화",
            "  normal = 거울 transform.up 정규화",
            "  dot = dir . normal",
            "  reflect = dir - 2 * dot * normal",
            "",
            "광선 추적:",
            "  RaycastAll 중 ignoreCollider와 매우 가까운 hit를 제외",
            "  Target이면 타깃 clear",
            "  Boundary이면 추적 종료",
            "  Mirror이면 reflect 방향으로 다음 segment 진행",
            "  Prism이면 Red/Green/Blue 광선을 새 RayLine으로 추가",
        ],
    )
    add_table(
        doc,
        ["스테이지", "맵 크기", "거울", "타깃", "프리즘", "난이도 근거"],
        [
            ("1", "9 x 6", "2", "1", "0", "기본 반사와 단일 타깃"),
            ("2", "10 x 6", "4", "1", "0", "거울 수 증가와 시작 방향 변화"),
            ("3", "10 x 7", "3", "3", "1", "프리즘 분기와 다중 타깃 데이터 추가"),
        ],
        [850, 1050, 850, 850, 850, 4910],
        header_fill=LIGHT_GRAY,
        font_size=8.8,
    )
    add_table(
        doc,
        ["평가 관점", "설명"],
        [
            ("수학 개념", "벡터 정규화, 내적, 반사 공식, 방향 벡터 회전"),
            ("Unity 기술", "Physics2D.RaycastAll, Collider2D 필터링, LineRenderer 시각화"),
            ("안정성", "반사 반복을 최대 20회로 제한하여 무한 반사 루프를 방지"),
            ("보완 예정", "targetColor 데이터는 존재하나 현재 클리어 판정은 색상 일치 비교까지 확장 필요"),
        ],
        [1900, 7460],
        header_fill=LIGHT_BLUE,
        font_size=9.2,
    )


def add_data_design(doc):
    heading(doc, "8. 데이터 기반 스테이지 설계", 1)
    para(
        doc,
        "세 퍼즐 모두 ScriptableObject를 사용하여 규칙 데이터와 맵 데이터를 코드에서 분리합니다. "
        "따라서 새로운 스테이지를 추가할 때 핵심 로직을 수정하지 않고 asset 데이터만 확장하는 구조를 지향합니다.",
    )
    add_table(
        doc,
        ["데이터 클래스", "주요 필드", "사용 위치", "설계 의도"],
        [
            ("StageData", "stage, map, mapSize, start, end, fixedBlocks", "pipePuzzle_Manager", "그리드 기반 파이프 맵과 도착 조건 정의"),
            ("bitMaskData", "stage, switchType, switchInfo", "BitMaskPuzzle", "스위치 연산 타입과 5비트 마스크 패턴 정의"),
            ("MirrorStageData", "stage, map, targets, mapSize, rayStartPoint, rayStartDir", "MirrorManager", "광선 퍼즐 맵, 타깃 색상, 시작 위치/방향 정의"),
        ],
        [1900, 3150, 2000, 2310],
        header_fill=LIGHT_BLUE,
        font_size=8.7,
    )


def add_class_matrix(doc):
    heading(doc, "9. 클래스 구조 및 책임", 1)
    add_table(
        doc,
        ["클래스", "책임", "주요 협력 대상", "핵심 함수/필드"],
        [
            ("FlowManager", "전역 게임 흐름, 씬 전환, 스테이지/퍼즐 클리어 상태 관리", "SceneManager, TMP_Text, 각 퍼즐 매니저", "selectStage, GoPuzzle, RequestScene, Clear, WriteLog"),
            ("BitMaskPuzzle", "스위치 생성, 비트 상태 연산, 램프 UI 갱신, 클리어 판정", "bitMaskData, BitMaskPuzzle_switch, FlowManager", "curState, switchs, interSwitch, lampUpdate"),
            ("BitMaskPuzzle_switch", "개별 스위치 입력을 BitMaskPuzzle에 전달", "BitMaskPuzzle", "onoff, switchNum, type"),
            ("pipePuzzle_Manager", "파이프 맵 생성, DFS 도달성 탐색, 연결 상태 표시", "StageData, pipePuzzle_Block, FlowManager", "makeStage, setupBlock, clearCheck"),
            ("pipePuzzle_Block", "블록 회전, 연결 가능한 방향 계산, 도로 색상 표시", "pipePuzzle_Manager", "turn, GetCanGo, onoffRoad"),
            ("MirrorManager", "거울 맵 생성, 광선 리스트 관리, 타깃 클리어 확인", "MirrorStageData, RayLine, Target, FlowManager", "makeMap, newRay, DrawRay, checkClear"),
            ("Mirror", "마우스 입력에 따른 거울 회전과 반사 방향 계산", "MirrorManager, InputSystem", "Update, Trun, calculateReflexDgree"),
            ("RayLine", "Raycast 기반 광선 경로 계산과 LineRenderer 반영", "MirrorManager, Mirror, Target", "calculateWay, GetValidHit, DrawLine"),
            ("Target", "타깃 색상/클리어 상태 보관", "RayLine, MirrorManager", "targetColor, isClear"),
        ],
        [1550, 3000, 2500, 2310],
        header_fill=LIGHT_GRAY,
        font_size=7.8,
    )
    heading(doc, "9.1 클래스 의존 흐름", 2)
    code_block(
        doc,
        [
            "FlowManager",
            "  ├─ PipeScene    -> pipePuzzle_Manager -> pipePuzzle_Block -> DFS clearCheck",
            "  ├─ BitMaskScene -> BitMaskPuzzle -> BitMaskPuzzle_switch -> bit operation",
            "  └─ MirrorScene  -> MirrorManager -> RayLine / Mirror / Target -> ray tracing",
            "",
            "ScriptableObject Data",
            "  ├─ StageData",
            "  ├─ bitMaskData",
            "  └─ MirrorStageData",
        ],
    )


def add_debug_and_quality(doc):
    heading(doc, "10. 디버그 기능 및 제출 전 품질 점검", 1)
    add_table(
        doc,
        ["항목", "현재 구현", "제출 전 권장 조치"],
        [
            ("내부 데이터 출력", "FlowManager.WriteLog가 12줄 단위로 TMP_Text에 로그 출력", "스택/비트/광선 로그가 실제 UI에 연결되어 보이는지 확인"),
            ("레벨 강제 스킵", "RequestScene과 stage 선택 구조는 존재", "심사 요구에 맞춘 '다음 스테이지' 또는 '퍼즐 클리어 처리' 버튼 연결 확인"),
            ("스테이지 인덱스", "stage가 1부터 시작하지만 일부 배열은 0 기반", "IsStageClear[stage]는 stage-1 보정 또는 배열 크기 조정 필요"),
            ("거울 타깃 색상", "targets와 BeamColor 데이터는 존재", "RayLine에서 beamColor와 targetColor 일치 여부를 비교하도록 확장 권장"),
            ("Target 색상 초기화", "MirrorManager가 Instantiate 후 targetColor를 대입", "Awake 이후 색상 반영이 필요하면 별도 Init 메서드로 처리"),
            ("패키지/리소스 출처", "Unity 패키지와 TextMesh Pro 리소스 사용", "제출 압축본 내 외부 리소스 명세표 포함"),
        ],
        [1900, 3730, 3730],
        header_fill=LIGHT_BLUE,
        font_size=8.4,
    )


def add_external_resources(doc):
    heading(doc, "11. 외부 리소스 및 라이브러리 명세", 1)
    para(
        doc,
        "아래 표는 프로젝트 파일 기준으로 확인된 외부 패키지와 리소스입니다. "
        "파이프 블록 PNG가 직접 제작물이 아니라 외부 제작물인 경우, 원 출처와 라이선스를 반드시 추가해야 합니다.",
        size=10.5,
    )
    add_table(
        doc,
        ["에셋/라이브러리명", "버전", "출처", "게임 내 구체적 활용도"],
        [
            ("Unity Editor", "6000.2.14f1", "Unity Technologies / https://unity.com", "게임 엔진 및 빌드 환경"),
            ("Universal Render Pipeline", "17.2.0", "Unity Package Manager / com.unity.render-pipelines.universal", "2D 렌더링 파이프라인"),
            ("Unity Input System", "1.16.0", "Unity Package Manager / com.unity.inputsystem", "마우스 좌/우 버튼 입력 처리"),
            ("Unity UGUI", "2.0.0", "Unity Package Manager / com.unity.ugui", "버튼, 이미지, 패널 등 UI 표시"),
            ("TextMesh Pro resources", "Unity bundled", "Unity Technologies / TextMesh Pro package resources", "디버그 텍스트 및 UI 글꼴 리소스"),
            ("pipePuzzle block PNG", "프로젝트 내 리소스", "확인 필요: 직접 제작 또는 원 출처 기재", "파이프 블록 종류별 이미지 표시"),
        ],
        [2050, 1200, 3050, 3060],
        header_fill=LIGHT_GRAY,
        font_size=7.8,
    )


def add_ai_usage(doc):
    heading(doc, "12. AI 활용 명세", 1)
    add_table(
        doc,
        ["사용 도구(모델)", "사용 목적", "구체적 활용 범위", "본인 검증/수정 여부"],
        [
            (
                "OpenAI Codex / GPT-5",
                "문서 작성 보조",
                "대회 안내서 요구사항 정리, Unity 코드 구조 분석 요약, 기술명세서 초안 작성 및 Word 문서 생성",
                "원본 프로젝트 코드와 안내서 PDF를 대조하여 반영. 핵심 퍼즐 알고리즘 코드는 본 작업에서 새로 생성하지 않음",
            ),
        ],
        [1900, 1800, 3400, 2260],
        header_fill=LIGHT_BLUE,
        font_size=8.0,
    )


def add_submission_checklist(doc):
    heading(doc, "13. 제출 패키지 체크리스트", 1)
    add_table(
        doc,
        ["제출물", "포함 권장 내용", "상태"],
        [
            ("프로그램 핵심 코드", "FlowManager.cs, 각 퍼즐 Manager/Block/Switch/RayLine/StageData 코드", "준비 가능"),
            ("프로그램 기술 명세서", "본 DOCX 파일", "작성 완료"),
            ("발표 자료", "Unity 버전, 타겟 플랫폼, 게임 구조, 문제 해결 과정, 핵심 설계 기술", "다음 작업 예정"),
            ("게임 플레이 영상", "시작, 플레이, 클리어/실패, 디버그 화면", "촬영 필요"),
            ("압축 파일명", "프로그래밍-성명-생년월일.zip", "성명/생년월일 확인 필요"),
        ],
        [2100, 5660, 1600],
        header_fill=LIGHT_GRAY,
        font_size=9.0,
    )


def add_appendix(doc):
    heading(doc, "부록 A. 핵심 코드 제출 권장 목록", 1)
    add_table(
        doc,
        ["분류", "파일"],
        [
            ("공통 흐름", "Assets/Script/FlowManager.cs"),
            ("비트마스크 퍼즐", "Assets/Script/BitMask/BitMaskPuzzle.cs, BitMaskPuzzle_switch.cs, Assets/Resources/bitMaskPuzzle/stage/bitMaskData.cs"),
            ("파이프 퍼즐", "Assets/Script/pipePuzzle/puzzleManager.cs, puzzleBlock.cs, Assets/Resources/pipePuzzle/stage/StageData.cs"),
            ("거울 퍼즐", "Assets/Script/mirrorPuzzle/MirrorManager.cs, Mirror.cs, RayLine.cs, Target.cs, Assets/Resources/Mirror/stage/MirrorStageData.cs"),
            ("스테이지 데이터", "Assets/Resources/bitMaskPuzzle/stage/*.asset, Assets/Resources/pipePuzzle/stage/*.asset, Assets/Resources/Mirror/stage/*.asset"),
        ],
        [2100, 7260],
        header_fill=LIGHT_BLUE,
        font_size=8.6,
    )
    heading(doc, "부록 B. 설계 프리셋", 1)
    para(
        doc,
        "문서는 standard_business_brief 계열을 기준으로 작성하되, 한국어 가독성을 위해 기본 글꼴은 맑은 고딕으로 조정했습니다. "
        "표는 고정 폭 DXA 기준으로 생성하여 Word 렌더링 시 폭이 흔들리지 않도록 구성했습니다.",
        size=10.5,
    )


def build():
    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_requirement_mapping(doc)
    add_overview(doc)
    add_environment(doc)
    add_scene_flow(doc)
    add_bitmask(doc)
    add_pipe(doc)
    add_mirror(doc)
    add_data_design(doc)
    add_class_matrix(doc)
    add_debug_and_quality(doc)
    add_external_resources(doc)
    add_ai_usage(doc)
    add_submission_checklist(doc)
    add_appendix(doc)

    doc.core_properties.title = "ChungKang_subject 프로그램 기술 명세서"
    doc.core_properties.subject = "2026 청강게임대전 게임프로그래밍"
    doc.core_properties.author = "Codex document assistant"
    doc.core_properties.keywords = "Unity, puzzle, technical specification, ChungKang_subject"
    doc.core_properties.created = datetime(2026, 7, 9, 0, 0, 0)
    doc.core_properties.modified = datetime(2026, 7, 9, 0, 0, 0)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
